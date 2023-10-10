

from docx import Document
import pandas as pd
from nltk.corpus import stopwords

from sklearn.neighbors import NearestNeighbors

from resume_parser import ResumeParser

import os

import re

from ftfy import fix_text

from sklearn.feature_extraction.text import TfidfVectorizer
import re

filed="file_path"

try:
    # doc = Document()
    # with open(filed, 'r') as file:
    #     doc.add_paragraph(file.read())
    # doc.save("text.docx")
    # data = ResumeParser('text.docx').get_extracted_data()
    # print(data['skills'])
    print("HI")
except:
    # data = ResumeParser(filed).get_extracted_data()
    # print(data['skills'])
    print("HI")

# stopw  = set(stopwords.words('english'))

# df =pd.read_csv('job_final.csv') 

# df['test']=df['Job_Description'].apply(lambda x: ' '.join([word for word in str(x).split() if len(word)>2 and word not in (stopw)]))
# df['test']

# ##file format should be in .txt , .docx or .pdf only
# filed=input('specify the path of the resume (format(.txt, .docx and .pdf))==')

# try:
#     doc = Document()
#     with open(filed, 'r') as file:
#         doc.add_paragraph(file.read())
#     doc.save("text.docx")
#     data1 = ResumeParser('text.docx').get_extracted_data()

# except:
#     data1 = ResumeParser(filed).get_extracted_data()
    
# resume=data1['skills']

# skills=[]
# skills.append(' '.join(word for word in resume))



# def ngrams(string, n=3):
#     string = fix_text(string) # fix text
#     string = string.encode("ascii", errors="ignore").decode() #remove non ascii chars
#     string = string.lower()
#     chars_to_remove = [")","(",".","|","[","]","{","}","'"]
#     rx = '[' + re.escape(''.join(chars_to_remove)) + ']'
#     string = re.sub(rx, '', string)
#     string = string.replace('&', 'and')
#     string = string.replace(',', ' ')
#     string = string.replace('-', ' ')
#     string = string.title() # normalise case - capital at start of each word
#     string = re.sub(' +',' ',string).strip() # get rid of multiple spaces and replace with a single
#     string = ' '+ string +' ' # pad names for ngrams...
#     string = re.sub(r'[,-./]|\sBD',r'', string)
#     ngrams = zip(*[string[i:] for i in range(n)])
#     return [''.join(ngram) for ngram in ngrams]





# vectorizer = TfidfVectorizer(min_df=1, analyzer=ngrams, lowercase=False)
# tfidf = vectorizer.fit_transform(skills)

# nbrs = NearestNeighbors(n_neighbors=1, n_jobs=-1).fit(tfidf)
# test = (df['test'].values.astype('U'))

# def getNearestN(query):
#   queryTFIDF_ = vectorizer.transform(query)
#   distances, indices = nbrs.kneighbors(queryTFIDF_)
#   return distances, indices

# distances, indices = getNearestN(test)
# test = list(test) 
# matches = []

# for i,j in enumerate(indices):
#     dist=round(distances[i][0],2)
  
#     temp = [dist]
#     matches.append(temp)
    
# matches = pd.DataFrame(matches, columns=['Match confidence'])

# df['match']=matches['Match confidence']
# df1=df.sort_values('match')
# df1[['Position', 'Company','Location']].head(10).reset_index()
